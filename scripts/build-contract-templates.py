"""
Gera os 4 templates DOCX de contrato para a BoomLab.

Placeholders sao os do docxtemplater: {variavel} (chavetas simples).

Variantes:
  A = 1 outorgante + prestacoes
  B = 2 outorgantes + prestacoes
  C = 1 outorgante + a vista
  D = 2 outorgantes + a vista

Output: play-store-assets/contract-templates/*.docx
Depois o utilizador sobe via /admin/contract-templates.
"""

import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "..", "contract-templates-output")
os.makedirs(OUTPUT_DIR, exist_ok=True)


BOOMLAB_HEADER = (
    "Boomlab Agency OÜ, com sede na Lasnamäe linnaosa, Ruunaoja tn 3, 11415 com o numero unico de "
    "registo 17418850, com o numero unico de pessoa coletiva, matricula e contribuinte fiscal "
    "EE102949302, aqui representada pelos seu gerentes Martim Salvador Fernandes Francisco, portador "
    "do Cartao de Cidadao n.º 30907570 0 ZW0, emitido pela Republica Portuguesa, valido ate "
    "07.08.2028 e Guilherme Cacais de Barros Freitas, portador do Cartao de Cidadao n.º 30297087 8 "
    "ZW5, emitido pela Republica Portuguesa, valido ate 09.08.2026, ambos com qualidade e "
    "suficiencia de poderes para o ato, adiante designada como PRIMEIRA OUTORGANTE,"
)

SEGUNDA_1_OUT = (
    "{nome_empresa}, com sede em {sede_empresa}, com o numero unico de registo de pessoa coletiva, "
    "matricula e contribuinte fiscal {nif_empresa}, aqui representada pelo/a gerente {gerente_1_nome} "
    "portador/a do Cartao de Cidadao n.º {gerente_1_cc}, emitido pela Republica Portuguesa, valido "
    "ate {gerente_1_cc_validade}, adiante designada como SEGUNDA OUTORGANTE,"
)

SEGUNDA_2_OUT = (
    "{nome_empresa}, com sede em {sede_empresa}, com o numero unico de registo de pessoa coletiva, "
    "matricula e contribuinte fiscal {nif_empresa}, aqui representada pelo/a gerente {gerente_1_nome} "
    "portador/a do Cartao de Cidadao n.º {gerente_1_cc}, emitido pela Republica Portuguesa, valido "
    "ate {gerente_1_cc_validade}, e pelo/a gerente {gerente_2_nome}, portador/a do Cartao de Cidadao "
    "n.º {gerente_2_cc}, emitido pela Republica Portuguesa, valido ate {gerente_2_cc_validade}, "
    "adiante designada como SEGUNDA OUTORGANTE,"
)

INTRO = (
    "E livremente e de boa-fe celebrado e reciprocamente aceite, o contrato de prestacao de servicos, "
    "celebrado ao abrigo do disposto nos artigos 1154º e seguintes do Codigo Civil que se rege nos "
    "termos e pelas seguintes clausulas:"
)

OBJETO_ITEMS = [
    "Criacao de uma nova oferta ou reestruturacao da oferta existente",
    "Implementacao de sistemas e processos que visam garantir a longevidade dos clientes angariados durante o periodo do contrato",
    "Apoio na organizacao de campanhas de prospecao de mercado atraves de meios como e-mails, chamadas telefonicas, e trafego pago nas redes sociais.",
    "Reunioes semanais de comunicacao de performance com o objetivo de analisar o sucesso do trabalho praticado.",
    "Implementacao ou melhoria do processo de integracao de novos clientes e membros de equipa",
    "Otimizacao dos sistemas de comunicacao internos e externos",
    "Apoio e orientacao na pesquisa, recrutamento e manutencao de novos membros de equipa",
    "Formacao continua em metodos de prospecao e tecnicas de venda",
    "Organizacao da estrutura e otimizacao dos processos de gestao interna",
]

DURACAO = (
    "O presente contrato de prestacao de servicos e celebrado pelo prazo de 4 meses, com inicio a "
    "{data_inicio} e termo a {data_fim}."
)

REMUN_PRESTACOES = [
    "A SEGUNDA OUTORGANTE paga na presente data a PRIMEIRA OUTORGANTE, a quantia de {primeira_prestacao}, atraves de deposito ou transferencia bancaria para a conta com o IBAN LT793250067540618109.",
    "A titulo de remuneracao pelos servicos prestados, a SEGUNDA OUTORGANTE pagara mensalmente a PRIMEIRA OUTORGANTE a quantia de {restantes_prestacoes}, atraves de deposito ou transferencia bancaria para a conta com o IBAN LT79 3250 0675 4061 8109.",
]

REMUN_AVISTA = [
    "A SEGUNDA OUTORGANTE paga na presente data a PRIMEIRA OUTORGANTE, a quantia de {primeira_prestacao}, atraves de deposito ou transferencia bancaria para a conta com o IBAN LT793250067540618109.",
]

RESULTADOS_PRESTACOES = [
    "A SEGUNDA OUTORGANTE podera resolver o presente contrato no prazo de 30 (trinta) dias a contar da data do seu inicio, mediante comunicacao escrita dirigida a PRIMEIRA OUTORGANTE.",
    "Em caso de resolucao dentro do prazo referido no numero anterior, a PRIMEIRA OUTORGANTE compromete-se a devolver a SEGUNDA OUTORGANTE a quantia correspondente a primeira prestacao paga, referida na clausula 3.ª, no prazo maximo de 5 (cinco) dias uteis apos a rececao da respetiva comunicacao.",
]

RESULTADOS_AVISTA = [
    "No 1º mes, caso a maneira de operar da PRIMEIRA OUTORGANTE nao corresponda as expectativas da SEGUNDA OUTORGANTE, a PRIMEIRA OUTORGANTE compromete-se a devolver a quantia referida no ponto 1 da 3a clausula, que a SEGUNDA OUTORGANTE pagou aquando do inicio do contrato.",
]

INCUMPRIMENTO = [
    "A falta de pagamento da remuneracao estipulada na clausula terceira provoca a suspensao imediata dos servicos prestados e todas as cominacoes legalmente estabelecidas respeitantes a incumprimentos contratuais.",
    "O recibo da prestacao de servicos relativo a cada mes sera emitido apos boa cobranca do pagamento da remuneracao acima acordada.",
]

CONFIDENCIALIDADE = [
    "Na vigencia do presente contrato e apos o seu termo, a PRIMEIRA OUTORGANTE obriga-se, em absoluto, a manter sigilo e confidencialidade sobre quaisquer informacoes ou dados que se refiram a SEGUNDA OUTORGANTE, aos seus clientes, negocios e ativos, incluindo, mas nao restrito, informacoes economicas, financeiras, contabilisticas e administrativas, especificacoes tecnicas, segredos comerciais, processos de comercializacao, formulas, know-how, estrategias, projetos, precos e margens de lucro.",
    "A violacao das obrigacoes assumidas no numero anterior constitui a PRIMEIRA OUTORGANTE na obrigacao de indemnizar a SEGUNDA OUTORGANTE, nos termos gerais de direito.",
]

AUTONOMIA = [
    "A PRIMEIRA OUTORGANTE presta os servicos de forma autonoma e independente, sem qualquer tipo de subordinacao juridica em relacao a SEGUNDA OUTORGANTE.",
    "Em virtude da posicao de autonomia e independencia da PRIMEIRA OUTORGANTE, esta e a unica responsavel pelo pagamento de todos os impostos, contribuicoes e quotizacoes de Seguranca Social, taxas, seguros e demais encargos inerentes a sua atividade de profissional independente.",
]

RESPONSAB = [
    "A PRIMEIRA OUTORGANTE prestara os servicos ora contratados com zelo, dedicacao e diligencia e em estreita colaboracao com a SEGUNDA OUTORGANTE.",
    "As partes contratantes sabem, reconhecem e aceitam que os servicos ora contratados sao prestados pela PRIMEIRA OUTORGANTE sem qualquer garantia de adequacao a determinado resultado para alem do referido na 4ª clausula, pelo que a PRIMEIRA OUTORGANTE nao sera responsavel perante a SEGUNDA OUTORGANTE ou terceiros, por quaisquer danos, incluindo lucros cessantes, decorrentes do atraso ou frustracao do efeito util pretendido com o presente contrato nem qualquer tipo de inadequacao.",
]

PROP_INTELECTUAL_PRESTACOES = [
    "No ambito da execucao do presente contrato, todos os dados, documentos, materiais, processos, estrategias, sistemas, conteudos, apresentacoes, relatorios ou quaisquer outros elementos desenvolvidos ou criados especificamente para a SEGUNDA OUTORGANTE poderao ser utilizados pela SEGUNDA OUTORGANTE apos o termo do presente contrato.",
    "Para efeitos do disposto no numero anterior, a PRIMEIRA OUTORGANTE concede a SEGUNDA OUTORGANTE uma licenca de utilizacao nao exclusiva, sem limitacao temporal, para uso no ambito da sua atividade.",
    "Permanecem propriedade da PRIMEIRA OUTORGANTE os metodos, know-how, frameworks, metodologias e processos de trabalho de carater geral, podendo estes ser livremente utilizados pela PRIMEIRA OUTORGANTE em outros projetos ou com terceiros.",
]

PROP_INTELECTUAL_AVISTA = [
    "A titularidade plena e exclusiva dos respetivos direitos de propriedade intelectual pertence a PRIMEIRA OUTORGANTE, nao ocorrendo, em momento algum, a sua transmissao ou cessao para a SEGUNDA OUTORGANTE.",
    "A SEGUNDA OUTORGANTE apenas podera utilizar os referidos resultados nos termos e para os fins expressamente previstos no presente contrato, nao podendo reproduzi-los, modifica-los, ceda-los a terceiros ou explora-los comercialmente sem autorizacao previa e escrita da PRIMEIRA OUTORGANTE.",
]

ALTERACOES = (
    "Quaisquer alteracoes ao presente contrato so serao validas desde que convencionadas por escrito, "
    "com a mencao expressa de cada uma das clausulas eliminadas e da redacao que passa a ter cada "
    "uma das aditadas ou alteradas."
)

INVALIDADE = [
    "A nulidade ou anulacao parcial nao determina a invalidade de todo o contrato.",
    "Em caso de invalidade de alguma das clausulas do presente contrato, as OUTORGANTES comprometem-se a proceder a respetiva substituicao por uma clausula valida, desde que tal substituicao seja justificada pelo interesse das partes e o objeto do contrato.",
]

FORO = (
    "O presente contrato rege-se pela lei portuguesa. Em tudo o que o presente contrato for omisso, "
    "regera a legislacao aplicavel, sendo para qualquer litigio ou questao de interpretacao "
    "exclusivamente competente o foro da comarca do Seixal, com expressa renuncia a qualquer outro."
)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)


def add_clause_title(doc, numero, titulo):
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.add_run(numero).bold = True
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run(f"({titulo})").bold = True


def add_list(doc, items, numbered=True):
    for idx, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        if numbered:
            p.add_run(f"{idx}. {item}")
        else:
            p.add_run(f"\u2022 {item}")


def build_contract(output_filename, *, two_outorgantes, prestacoes):
    doc = Document()

    # Styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    add_heading(doc, "CONTRATO DE PRESTACAO DE SERVICOS", level=0)
    doc.add_paragraph()
    doc.add_paragraph("ENTRE,")

    # First outorgante
    doc.add_paragraph(BOOMLAB_HEADER)
    doc.add_paragraph("E")

    # Second outorgante
    if two_outorgantes:
        doc.add_paragraph(SEGUNDA_2_OUT)
    else:
        doc.add_paragraph(SEGUNDA_1_OUT)

    doc.add_paragraph()
    doc.add_paragraph(INTRO)
    doc.add_paragraph()

    # 1 - Objeto
    add_clause_title(doc, "1ª", "Objeto")
    doc.add_paragraph(
        "O presente contrato tem por objeto a prestacao de servicos de consultoria comercial e "
        "estrategica a prestar pela PRIMEIRA OUTORGANTE a SEGUNDA OUTORGANTE, nomeadamente:"
    )
    add_list(doc, OBJETO_ITEMS, numbered=True)
    doc.add_paragraph()

    # 2 - Duracao
    add_clause_title(doc, "2ª", "Duracao")
    doc.add_paragraph(DURACAO)
    doc.add_paragraph()

    # 3 - Remuneracao
    add_clause_title(doc, "3ª", "Remuneracao")
    remun = REMUN_PRESTACOES if prestacoes else REMUN_AVISTA
    add_list(doc, remun, numbered=True)
    doc.add_paragraph()

    # 4 - Resultados
    add_clause_title(doc, "4ª", "Resultados")
    resultados = RESULTADOS_PRESTACOES if prestacoes else RESULTADOS_AVISTA
    add_list(doc, resultados, numbered=True)
    doc.add_paragraph()

    # 5 - Incumprimento
    add_clause_title(doc, "5ª", "Incumprimento")
    add_list(doc, INCUMPRIMENTO, numbered=True)
    doc.add_paragraph()

    # 6 - Confidencialidade
    add_clause_title(doc, "6ª", "Confidencialidade")
    add_list(doc, CONFIDENCIALIDADE, numbered=True)
    doc.add_paragraph()

    # 7 - Autonomia
    add_clause_title(doc, "7ª", "Autonomia e independencia")
    add_list(doc, AUTONOMIA, numbered=True)
    doc.add_paragraph()

    # 8 - Exclusao de responsabilidade
    add_clause_title(doc, "8ª", "Exclusao da responsabilidade")
    add_list(doc, RESPONSAB, numbered=True)
    doc.add_paragraph()

    # 9 - Propriedade intelectual
    add_clause_title(doc, "9ª", "Propriedade intelectual")
    propint = PROP_INTELECTUAL_PRESTACOES if prestacoes else PROP_INTELECTUAL_AVISTA
    add_list(doc, propint, numbered=True)
    doc.add_paragraph()

    # 10 - Alteracoes
    add_clause_title(doc, "10ª", "Alteracoes")
    doc.add_paragraph(ALTERACOES)
    doc.add_paragraph()

    # 11 - Invalidade parcial
    add_clause_title(doc, "11ª", "Invalidade parcial do contrato")
    add_list(doc, INVALIDADE, numbered=True)
    doc.add_paragraph()

    # 12 - Foro
    add_clause_title(doc, "12ª", "Foro")
    doc.add_paragraph(FORO)
    doc.add_paragraph()
    doc.add_paragraph()

    # Signatures
    p = doc.add_paragraph()
    p.add_run("Seixal, {data_assinatura}")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("PRIMEIRA OUTORGANTE:")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("_________________________________")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("SEGUNDA OUTORGANTE:")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("_________________________________")
    if two_outorgantes:
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph("_________________________________")

    path = os.path.join(OUTPUT_DIR, output_filename)
    doc.save(path)
    print(f"[OK] {path}")


build_contract("template-A-1outorgante-prestacoes.docx", two_outorgantes=False, prestacoes=True)
build_contract("template-B-2outorgantes-prestacoes.docx", two_outorgantes=True, prestacoes=True)
build_contract("template-C-1outorgante-avista.docx", two_outorgantes=False, prestacoes=False)
build_contract("template-D-2outorgantes-avista.docx", two_outorgantes=True, prestacoes=False)

print("\nTemplates gerados em:", os.path.abspath(OUTPUT_DIR))
